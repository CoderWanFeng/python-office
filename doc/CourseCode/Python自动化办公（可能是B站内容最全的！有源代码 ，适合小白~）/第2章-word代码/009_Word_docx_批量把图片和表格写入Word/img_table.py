# -*- coding: utf-8 -*-
# @Time : 2020/8/20 18:49
# @公众号 :Python自动化办公社区 
# @File : img_table.py
# @Software: PyCharm
# @Description:


from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt
from docx.oxml.ns import qn
from docx.shared import Inches

import time

price = input('请输入工资调整金额：')
company_list = ['员工1', '员工1', '员工2', '员工3', '员工4', '员工5', '员工6', '员工7', '员工8', '员工9', '员工10', ]
today = time.strftime("%Y{y}%m{m}%d{d}", time.localtime()).format(y='年', m='月', d='日')
for i in company_list:
    document = Document()
    # 设置文档的基础字体
    document.styles['Normal'].font.name = u'宋体'
    document.styles['Normal'].element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
    # 红头文件
    document.add_picture('title002.jpg', width=Inches(6))

    # 建立一个自然段
    p1 = document.add_paragraph()
    # 对齐方式为居中，没有这句的话默认左对齐
    p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run1 = p1.add_run('关于%s工资调整的通知' % (today))
    run1.font.name = '微软雅黑'
    run1.element.rPr.rFonts.set(qn('w:eastAsia'), u'微软雅黑')
    run1.font.size = Pt(21)
    run1.font.bold = True
    p1.space_after = Pt(5)
    p1.space_before = Pt(5)

    p2 = document.add_paragraph()
    run2 = p2.add_run(i + '：')
    run2.font.name = '宋体'
    run2.element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
    run2.font.size = Pt(16)
    run2.font.bold = True

    p3 = document.add_paragraph()
    run3 = p3.add_run('因为疫情影响，我们很抱歉的通知您，您的工资调整为每月%s元，特此通知。' % price)
    run3.font.name = '宋体'
    run3.element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
    run3.font.size = Pt(14)

    table = document.add_table(rows=2, cols=2, style='Table Grid')
    # 合并单元格
    table.cell(0, 0).merge(table.cell(0, 1))
    table_run1 = table.cell(0, 0).paragraphs[0].add_run('签名栏')
    table_run1.font.name = '黑体'
    table_run1.element.rPr.rFonts.set(qn('w:eastAsia'), u'黑体')
    table.cell(0, 0).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    table.cell(1, 0).text = i
    table.cell(1, 0).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    p4 = document.add_paragraph()
    p4.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    run4 = p4.add_run('人事：王小姐 电话：686868')
    run4.font.name = '宋体'
    run4.element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
    run4.font.size = Pt(14)
    run4.font.bold = True

    document.save('%s-工资调整通知.docx' % i)
